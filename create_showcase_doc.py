#!/usr/bin/env python3
"""
Create a professional Word document showcasing the fhir-query-validator-factory project.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_mermaid_block(doc, code):
    """Add a Mermaid diagram as a formatted code block."""
    para = doc.add_paragraph()
    run = para.add_run("Mermaid Diagram (copy to https://mermaid.live to view):")
    run.bold = True
    run.font.size = Pt(10)
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(8)

def create_document():
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.author = "Yogesh"
    core_properties.title = "fhir-query-validator-factory - Software Factory Demonstration"
    
    # Title
    title = doc.add_heading('fhir-query-validator-factory', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('A Software Factory Demonstration for Modern Agentic AI Systems')
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ============================================
    # 1. INTRODUCTION
    # ============================================
    doc.add_heading('1. Introduction', level=1)
    
    intro_text = """This document presents a comprehensive demonstration of applying Software Factory principles to the development of agentic AI systems. The project builds upon the foundational work done in the original fhirqueryvalidator repository (https://github.com/yogesh-parte/fhirqueryvalidator), which implemented a focused, rule-based pre-flight validator for FHIR search queries.

While the original repository provided a solid, practical tool for validating FHIR queries against a server's CapabilityStatement, this new demonstration takes a fundamentally different approach. Instead of building another validator in the traditional way, we have created fhir-query-validator-factory as a reference implementation that showcases how modern software engineering practices — specifically the Software Factory pattern combined with loop engineering and spec-driven development — can be applied to agentic AI systems.

The core objective is to demonstrate that building reliable, governable, and observable agentic systems requires the same engineering discipline that has proven successful in traditional software development: clear specifications, phased planning, modular design, explicit feedback loops, and human oversight at critical points."""
    
    doc.add_paragraph(intro_text)
    
    # ============================================
    # 2. BUSINESS OBJECTIVE
    # ============================================
    doc.add_heading('2. Business Objective', level=1)
    
    obj_text = """The primary business objective of this project is to establish a repeatable, disciplined methodology for developing agentic AI systems that can be trusted in production environments, particularly in regulated domains such as healthcare.

Traditional approaches to building agentic systems often suffer from:
• Ad-hoc development with insufficient upfront planning
• Monolithic agent designs that are difficult to maintain and debug
• Lack of explicit feedback loops, leading to silent failures or repeated errors
• Insufficient human oversight in critical decision paths
• Poor traceability, making it difficult to understand why an agent made a particular decision

This demonstration addresses these challenges by showing how Software Factory principles can be adapted for the age of agentic AI. The specific technical objective was to evolve a conventional FHIR query validator into a more intelligent, generalized system that can:
• Dynamically validate any query parameter declared in a CapabilityStatement (not just pre-defined ones)
• Execute valid queries against FHIR servers (including authenticated ones)
• Detect repeated invalid query patterns from users
• Provide automated guidance or escalate to human review when appropriate

By achieving these technical goals through a structured process, the project serves as a blueprint for future agentic initiatives within the organization."""
    
    doc.add_paragraph(obj_text)
    
    # ============================================
    # 3. VALUES DELIVERED
    # ============================================
    doc.add_heading('3. Values Delivered', level=1)
    
    values_intro = doc.add_paragraph("This project delivers value across multiple dimensions:")
    
    doc.add_heading('3.1 Engineering Discipline', level=2)
    doc.add_paragraph("By following a phased planning approach (detailed in the planning/ folder) and creating comprehensive specifications before implementation, the project demonstrates that even agentic systems benefit from traditional software engineering rigor. This reduces rework and improves maintainability.")
    
    doc.add_heading('3.2 Repeatable Pattern', level=2)
    doc.add_paragraph("The combination of planning artifacts, specifications, architecture decisions (documented in ADRs), and implementation provides a reusable template that can be applied to other agentic use cases beyond FHIR query validation.")
    
    doc.add_heading('3.3 Responsible AI', level=2)
    doc.add_paragraph("The explicit design of feedback loops — particularly the meta-loop involving pattern detection, automated guidance, and human escalation — demonstrates how agentic systems can be made safer and more aligned with human intent, which is critical in healthcare contexts.")
    
    doc.add_heading('3.4 Observability & Governance', level=2)
    doc.add_paragraph("Through structured logging, traceability reports, and optional integration with observability platforms like Langfuse, the system makes agent decisions visible and auditable — a key requirement for governance in regulated industries.")
    
    doc.add_heading('3.5 Knowledge Management', level=2)
    doc.add_paragraph("The emphasis on specifications, planning documents, and the recommendation to use tools like the OKF (Open Knowledge Format) skill ensures that institutional knowledge is captured and can be easily consumed by future team members or leadership.")
    
    # ============================================
    # 4. WHAT DOES IT DO?
    # ============================================
    doc.add_heading('4. What Does It Do?', level=1)
    
    doc.add_paragraph("At its core, fhir-query-validator-factory implements a generalized, intelligent FHIR query validation and execution system with the following capabilities:")
    
    doc.add_heading('4.1 Generalized Validation', level=2)
    doc.add_paragraph("Unlike the original fhirqueryvalidator which had limited, pre-defined validation rules, this system dynamically interprets a FHIR server's CapabilityStatement and can validate any search parameter, modifier, or comparator that the server declares as supported.")
    
    doc.add_heading('4.2 Query Execution', level=2)
    doc.add_paragraph("When configured in 'validate_and_execute' mode, the system not only validates queries but also executes them against the target FHIR server (supporting both public test servers and authenticated servers via Bearer tokens or OAuth2).")
    
    doc.add_heading('4.3 Intelligent Feedback Loops', level=2)
    doc.add_paragraph("The system implements multiple explicit feedback loops:")
    
    bullets = [
        "Cache Invalidation Loop: Efficiently manages CapabilityStatement freshness using TTL and conditional requests.",
        "Validation → Execution Loop: Ensures queries are only executed after successful validation.",
        "Pattern Detection → Learning / Human Escalation Loop (Meta Loop): Detects when a user repeatedly submits invalid queries and responds by either providing automated guidance (via the Search Learner Agent) or escalating to human review."
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_heading('4.4 Workflow Visualization', level=2)
    
    doc.add_paragraph("The following Mermaid diagram illustrates the main validation workflow and the feedback loops:")
    
    mermaid_code = """flowchart TD
    A[Start: Receive Query + server_key + user_id] --> B[CacheAgent: Get CapabilityStatement]
    B --> C{Cache Valid?}
    C -->|No| D[Fetch from Server]
    C -->|Yes| E[CapabilityInterpreter: Parse & Interpret]
    D --> E
    E --> F[QueryValidator: Validate Query + Detect Patterns]
    F --> G{Valid?}
    G -->|Yes + Execute Mode| H[QueryExecution: Execute against FHIR Server]
    G -->|No| I{Pattern Detected?}
    H --> J[Return Results]
    I -->|Yes| K[RuleAgent: Decide Escalation]
    I -->|No| L[Return Validation Errors]
    K --> M{Learner or Human?}
    M -->|Learner| N[SearchLearnerAgent: Provide Guidance]
    M -->|Human| O[HumanInterventionGate: Escalate]
    N --> P[Return Guidance + Errors]
    O --> Q[Pause & Notify Human]
    J --> R[End]
    L --> R
    P --> R
    Q --> R"""
    
    add_mermaid_block(doc, mermaid_code)
    
    doc.add_paragraph()
    doc.add_paragraph("The workflow clearly shows the separation between normal operational paths and the intelligent feedback paths that make the system more than just a simple validator.")
    
    # ============================================
    # 5. WHAT IT DOESN'T DO (OUT OF SCOPE)
    # ============================================
    doc.add_heading('5. What It Does Not Do (Out of Scope)', level=1)
    
    scope_text = """It is important to understand the boundaries of this demonstration:

• Production Readiness: This is a demonstration, not a production-ready library. It lacks comprehensive error handling, security hardening, rate limiting, and performance optimizations required for high-volume production use.

• Full FHIR Compliance: The validation logic is simplified for demonstration purposes. A production system would require deeper analysis of FHIR search semantics, value set validation, and complex parameter interactions.

• Persistent Learning: The pattern detection and learning mechanisms are in-memory only. A production system would likely require persistent storage and more sophisticated machine learning models for user behavior analysis.

• Advanced OAuth2 Flows: While basic Bearer token and OAuth2 client credentials are supported in the configuration, advanced flows (e.g., authorization code with PKCE, token refresh with rotation) are not fully implemented.

• Multi-tenancy & Security: The system does not implement tenant isolation, fine-grained access control, or audit logging suitable for multi-tenant SaaS environments.

• Real-time Human-in-the-Loop UI: The Human Intervention Gate is implemented as a logging/notification mechanism. A real production system would require integration with ticketing systems, approval workflows, and a user interface for reviewers.

This project intentionally focuses on demonstrating the process and architectural patterns rather than delivering a complete, deployable product."""
    
    doc.add_paragraph(scope_text)
    
    # ============================================
    # 6. CODE STRUCTURE
    # ============================================
    doc.add_heading('6. Code Structure & How to Use', level=1)
    
    doc.add_heading('6.1 Folder Structure', level=2)
    
    structure = """fhir-query-validator-factory/
├── docs/
│   ├── spec/                    # Agent specifications (source of truth)
│   ├── architecture.md
│   ├── loop-engineering.md
│   ├── traceability.md
│   ├── configuration.md
│   ├── process-overview.md
│   └── codex-context/           # Context package for LLM coding tools
├── planning/                    # Phase-by-phase planning artifacts
│   ├── phase-0-ideation.md
│   ├── phase-1-requirements-planning.md
│   ├── phase-2-architecture.md
│   ├── phase-3-scaffolding-core-agents.md
│   └── phase-4-loop-engineering.md
├── src/agentic_layer/
│   ├── agents/                  # All specialist agents
│   ├── config/                  # Configuration & server management
│   └── graph/                   # Main ADK workflow
├── scripts/
│   ├── demo_loops.py            # Demonstrates feedback loops
│   └── demo_traceability.py     # Structured traceability reports
├── examples/notebooks/          # Jupyter notebook demos
├── tests/
│   ├── unit/
│   ├── regression/
│   └── integration/
├── pyproject.toml
└── README.md"""
    
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(structure)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(8)
    
    doc.add_heading('6.2 How to Use', level=2)
    
    usage_text = """1. Configuration: Copy .env.example to .env.local and configure your servers and credentials.

2. Running Demos:
   - python scripts/demo_loops.py          (Shows feedback loops in action)
   - python scripts/demo_traceability.py   (Shows structured agent traces)

3. Running Tests:
   - pytest tests/ -v

4. Extending the System:
   - Add new specifications in docs/spec/
   - Implement new agents following the existing patterns
   - Wire new agents into the graph workflow in src/agentic_layer/graph/validation_workflow.py"""
    
    doc.add_paragraph(usage_text)
    
    # ============================================
    # 7. TESTING AND VALIDATION
    # ============================================
    doc.add_heading('7. Testing and Validation', level=1)
    
    doc.add_paragraph("The project includes a multi-layered testing strategy aligned with the Software Factory emphasis on quality gates:")
    
    doc.add_heading('7.1 Unit Tests', level=2)
    doc.add_paragraph("Located in tests/unit/. These tests verify individual agent behavior in isolation (e.g., CacheAgent cache hit/miss logic, QueryValidatorAgent pattern detection).")
    
    doc.add_heading('7.2 Regression Tests', level=2)
    doc.add_paragraph("Located in tests/regression/. These tests ensure that core validation behavior does not regress as the system evolves.")
    
    doc.add_heading('7.3 Integration Tests', level=2)
    doc.add_paragraph("Located in tests/integration/. These tests verify that the full workflow (multiple agents working together) produces expected outcomes.")
    
    doc.add_heading('7.4 How to Execute Tests', level=2)
    doc.add_paragraph("Run all tests: pytest tests/ -v")
    
    # ============================================
    # 8. LIMITATIONS
    # ============================================
    doc.add_heading('8. Limitations', level=1)
    
    limitations = """While this demonstration successfully showcases the Software Factory approach, several limitations should be acknowledged:

• Simplified Validation Logic: The current QueryValidatorAgent uses simplified logic for demonstration. A production implementation would require significantly more sophisticated FHIR search semantics parsing.

• In-Memory State: Pattern history and cache are stored in memory. Production deployments would require persistent, distributed storage.

• Demo-Grade Human-in-the-Loop: The Human Intervention Gate currently only logs the need for review. A real implementation would integrate with ticketing, notification, and approval systems.

• Limited Error Handling: The code prioritizes clarity and demonstration of concepts over exhaustive error handling and edge case coverage.

• Simulated Execution: The QueryExecutionAgent simulates FHIR server responses for many query patterns rather than always calling real servers.

These limitations are intentional. The primary goal is to demonstrate the process and architectural thinking, not to deliver production-ready components. Teams adopting this pattern should expect to invest additional effort in hardening, testing, and operationalizing the agents for their specific environment."""
    
    doc.add_paragraph(limitations)
    
    # ============================================
    # APPENDIX
    # ============================================
    doc.add_heading('Appendix A: Loop Engineering Deep Dive', level=1)
    
    appendix_a = """The concept of "Loop Engineering" is central to this demonstration. In traditional software, we often think in terms of linear request-response cycles. In agentic systems, we must deliberately design feedback loops that allow the system to observe its own behavior and adapt.

This project implements three categories of loops:

1. Operational Loops (Cache Invalidation, Validation → Execution): These improve efficiency and safety.
2. Learning Loops (Pattern Detection → Guidance): These allow the system to improve user experience over time without human intervention.
3. Governance Loops (Human Escalation): These ensure that when automation reaches its limits, humans remain in control.

The explicit design and documentation of these loops (see docs/loop-engineering.md) is what distinguishes a Software Factory approach from ad-hoc agent development."""
    
    doc.add_paragraph(appendix_a)
    
    doc.add_heading('Appendix B: The Planning Phases', level=1)
    
    appendix_b = """A core tenet of the Software Factory demonstrated here is that planning is not overhead — it is the highest-leverage activity.

• Phase 0 (Ideation): Defining the problem and vision.
• Phase 1 (Requirements): Translating vision into concrete requirements and success criteria.
• Phase 2 (Architecture): Making technology decisions, defining agent boundaries, and designing feedback loops.
• Phase 3 (Scaffolding & Core Agents): Implementing the foundation.
• Phase 4 (Loop Engineering): Completing the intelligent feedback mechanisms.

Detailed artifacts from each phase are preserved in the planning/ folder. This transparency allows others to understand not just what was built, but why and how the decisions were made."""
    
    doc.add_paragraph(appendix_b)
    
    doc.add_heading('Appendix C: Human-in-the-Loop Governance', level=1)
    
    appendix_c = """One of the most important aspects of responsible agentic system design is knowing when to stop automation and involve humans.

In this system, the Human Intervention Gate is triggered by the Rule Agent when repeated invalid patterns suggest that automated guidance is insufficient. This could indicate:
• A user struggling with the system interface
• A potential attempt to probe system boundaries
• A gap in the validation logic that requires human judgment

The design ensures that automation serves humans, rather than replacing human judgment in ambiguous or high-stakes situations. This is a key differentiator of the Software Factory approach demonstrated here."""
    
    doc.add_paragraph(appendix_c)
    
    # Final note
    doc.add_paragraph()
    final_note = doc.add_paragraph()
    final_note_run = final_note.add_run("This document and the accompanying repository serve as a reference example of how to approach the development of agentic AI systems with the rigor, transparency, and governance that modern engineering demands.")
    final_note_run.italic = True
    
    # Save the document
    output_path = "/home/workdir/artifacts/fhir-query-validator-factory/docs/fh ir-query-validator-factory-showcase.docx"
    # Fix filename (remove space)
    output_path = "/home/workdir/artifacts/fhir-query-validator-factory/docs/fhir-query-validator-factory-showcase.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
